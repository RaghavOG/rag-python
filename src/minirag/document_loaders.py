"""Document loaders: raw data → structured text + metadata."""
from pathlib import Path
from dataclasses import dataclass
from typing import Iterator

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None


@dataclass
class LoadedDocument:
    """Single document with content and metadata."""
    content: str
    source: str
    metadata: dict


def load_file(path: Path) -> LoadedDocument | None:
    """Load a single file (PDF, TXT, DOCX, MD) into text + metadata."""
    path = Path(path)
    if not path.exists():
        return None
    suffix = path.suffix.lower()
    metadata = {"source": str(path), "filename": path.name}

    if suffix == ".txt" or suffix == ".md":
        content = path.read_text(encoding="utf-8", errors="replace")
        return LoadedDocument(content=content, source=str(path), metadata=metadata)

    if suffix == ".pdf" and PdfReader:
        try:
            reader = PdfReader(path)
            parts = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                parts.append(text)
                metadata.setdefault("page_numbers", []).append(i + 1)
            content = "\n\n".join(parts)
            metadata["pages"] = len(parts)
            return LoadedDocument(content=content, source=str(path), metadata=metadata)
        except Exception:
            return None

    if suffix in (".docx", ".doc") and DocxDocument:
        try:
            doc = DocxDocument(path)
            parts = [p.text for p in doc.paragraphs]
            content = "\n\n".join(parts)
            metadata["paragraphs"] = len(parts)
            return LoadedDocument(content=content, source=str(path), metadata=metadata)
        except Exception:
            return None

    return None


def load_directory(dir_path: Path, extensions: tuple = (".txt", ".md", ".pdf", ".docx")) -> Iterator[LoadedDocument]:
    """Yield LoadedDocument for each supported file under dir_path."""
    dir_path = Path(dir_path)
    if not dir_path.is_dir():
        return
    for f in dir_path.rglob("*"):
        if f.is_file() and f.suffix.lower() in extensions:
            doc = load_file(f)
            if doc and doc.content.strip():
                yield doc

