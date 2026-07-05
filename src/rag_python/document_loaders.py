"""Document loaders: raw data → structured text + metadata."""
import csv
import json
from html.parser import HTMLParser
from pathlib import Path
from dataclasses import dataclass
from typing import Iterator

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None


@dataclass
class LoadedDocument:
    """Single document with content and metadata."""
    content: str
    source: str
    metadata: dict


class _HTMLTextExtractor(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []

    def handle_data(self, data: str) -> None:
        text = data.strip()
        if text:
            self.parts.append(text)


def _html_to_text(html: str) -> str:
    parser = _HTMLTextExtractor()
    parser.feed(html)
    return "\n".join(parser.parts)


def _load_csv(path: Path, metadata: dict) -> LoadedDocument | None:
    rows: list[str] = []
    with path.open(encoding="utf-8", errors="replace", newline="") as f:
        reader = csv.DictReader(f)
        if reader.fieldnames:
            for row in reader:
                rows.append(", ".join(f"{k}: {v}" for k, v in row.items() if v))
        else:
            f.seek(0)
            for row in csv.reader(f):
                rows.append(", ".join(row))
    content = "\n".join(rows)
    metadata["rows"] = len(rows)
    return LoadedDocument(content=content, source=str(path), metadata=metadata) if content.strip() else None


def _load_json(path: Path, metadata: dict) -> LoadedDocument | None:
    data = json.loads(path.read_text(encoding="utf-8", errors="replace"))
    if isinstance(data, list):
        parts = []
        for item in data:
            if isinstance(item, dict) and "text" in item:
                parts.append(str(item["text"]))
            else:
                parts.append(json.dumps(item, ensure_ascii=False))
        content = "\n\n".join(parts)
    elif isinstance(data, dict):
        if "text" in data:
            content = str(data["text"])
        else:
            content = json.dumps(data, ensure_ascii=False, indent=2)
    else:
        content = str(data)
    return LoadedDocument(content=content, source=str(path), metadata=metadata) if content.strip() else None


def load_file(path: Path) -> LoadedDocument | None:
    """Load a single file (PDF, TXT, DOCX, MD, CSV, JSON, HTML) into text + metadata."""
    path = Path(path)
    if not path.exists():
        return None
    suffix = path.suffix.lower()
    metadata = {"source": str(path), "filename": path.name}

    if suffix in (".txt", ".md"):
        content = path.read_text(encoding="utf-8", errors="replace")
        return LoadedDocument(content=content, source=str(path), metadata=metadata)

    if suffix == ".html":
        html = path.read_text(encoding="utf-8", errors="replace")
        content = _html_to_text(html)
        return LoadedDocument(content=content, source=str(path), metadata=metadata) if content.strip() else None

    if suffix == ".csv":
        return _load_csv(path, metadata)

    if suffix == ".json":
        try:
            return _load_json(path, metadata)
        except json.JSONDecodeError:
            return None

    if suffix == ".pdf" and PdfReader:
        try:
            reader = PdfReader(path)
            parts = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                parts.append(text)
                metadata.setdefault("page_numbers", []).append(i + 1)
            content = "\n\n".join(parts)
            metadata["pages"] = len(parts)
            return LoadedDocument(content=content, source=str(path), metadata=metadata)
        except Exception:
            return None

    if suffix in (".docx", ".doc") and DocxDocument:
        try:
            doc = DocxDocument(path)
            parts = [p.text for p in doc.paragraphs]
            content = "\n\n".join(parts)
            metadata["paragraphs"] = len(parts)
            return LoadedDocument(content=content, source=str(path), metadata=metadata)
        except Exception:
            return None

    return None


def load_directory(
    dir_path: Path,
    extensions: tuple = (".txt", ".md", ".pdf", ".docx", ".csv", ".json", ".html"),
) -> Iterator[LoadedDocument]:
    """Yield LoadedDocument for each supported file under dir_path."""
    dir_path = Path(dir_path)
    if not dir_path.is_dir():
        return
    for f in dir_path.rglob("*"):
        if f.is_file() and f.suffix.lower() in extensions:
            doc = load_file(f)
            if doc and doc.content.strip():
                yield doc
